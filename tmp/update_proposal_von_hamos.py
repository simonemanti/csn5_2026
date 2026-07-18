from pathlib import Path
from shutil import copy2

from docx import Document


ROOT = Path("/Users/smanti/WORK/BANDI/2026_07_CSN5")
SOURCE = ROOT / "Proposal_Grant_Giovani_2026.docx"
OUTPUT = ROOT / "Proposal_Grant_Giovani_2026_vonHamos_maxFlux.docx"


REPLACEMENTS = {
    "The project is divided into five Work Packages (WPs):": (
        "The project is divided into five Work Packages (WPs). The experimental architecture is a fixed constraint: "
        "a scan-free, energy-dispersive von Hamos geometry with a cylindrically bent HAPG crystal, the sample close "
        "to the X-ray source, and a position-sensitive CZT detector. This choice preserves continuity with the "
        "VOXES and operando laboratory-XAS work of the team; Johann/Johansson geometries are used only as external benchmarks."
    ),
    "This WP involves the development of simulations of the experimental setup in order to determine the most suitable components for the design phase of WP2.": (
        "WP1 will optimize the fixed von Hamos architecture rather than compare unrelated monochromator families. "
        "SHADOW4 ray tracing and GEANT4 radiation-transport simulations will maximize the useful photon rate accepted "
        "by the complete HAPG aperture and detected over a simultaneous bandwidth of approximately 0.5-1.5 keV. "
        "The optimization variables will include source-sample distance (baseline approximately 80 mm), HAPG radius, "
        "length and thickness, source take-off angle, filtration, and point/line-focus orientation. The short axis of "
        "a line focus will be aligned with the dispersive direction. Source ranking will use photon rate at the sample "
        "at 25.5, 40 and 50 keV, energy broadening, thermal stability and 8-24 h continuous-duty operation, not nominal kV/W alone."
    ),
    "Based on the results obtained in WP1 and within the allocated project budget, the most suitable crystals, motion stages, and CZT detectors will be selected and integrated into the laboratory high-energy XAS platform.": (
        "Based on WP1, WP2 will integrate the highest-flux tungsten-anode source compatible with the von Hamos etendue "
        "and the EUR 40k source-package cap. Two procurement tracks will be evaluated in parallel: (A) a 60 kV, 2 kW "
        "water-cooled diffraction tube with an apparent line focus for maximum flux at 20-30 keV; and (B) a 100 kV "
        "source delivering at least 500 W and preferably 1 kW continuously for extension toward 50 keV. Candidate RFQs "
        "include a KYW600-class W tube with a 60 kV/3.5 kW generator, and a Spellman DXM100-class generator with a "
        "water-cooled W tube having a small or line focus. An integrated 100 kV/500 W source is the lower-risk fallback. "
        "The existing 65 kV/100 W generator and 60 kV/75 W tube remain available for Phase-0 validation. HAPG optics, "
        "motion stages and CZT thickness, segmentation, rate capability, pile-up rejection and DAQ will be optimized concurrently."
    ),
    "To be defined.": (
        "The in-kind 60 kV/75 W source will first validate alignment, HAPG dispersion, energy calibration and CZT event "
        "reconstruction at the Mo K-edge. Measurements will quantify the photon rate accepted by the full optic, "
        "source-position stability, harmonic contamination, detector dead time and signal-to-background ratio before "
        "the high-power source is installed."
    ),
    "ddd": (
        "The source work package has a maximum allocation of EUR 40k. Suppliers will be asked to itemize the X-ray "
        "tube/housing (target allocation EUR 30k) and generator, cooling and control system (target allocation EUR 10k), "
        "or quote an integrated source within the same cap. Offers must state continuous current and power at 40, 50, "
        "60, 80 and 100 kV; tube hood, shutter, interlocks, cooling, CE documentation, delivery and acceptance testing. "
        "The final choice is conditional on competitive quotations and documented useful flux through the real HAPG geometry."
    ),
    "-Insufficient flux or signal-to-background ratio at the selected edges above 20 keV": (
        "- Insufficient useful flux or signal-to-background ratio above 20 keV. Mitigation: parallel RFQs for a 60 kV/2 kW "
        "line-focus source and a 100 kV/0.5-1.2 kW source; acceptance based on flux through the HAPG aperture at 25.5/40/50 keV; "
        "integrated 100 kV/500 W fallback; Phase-0 validation with the in-kind source; staged milestones, with EXAFS above the "
        "Ag K-edge retained as a stretch objective until the measured rate is sufficient.\n"
        "- Source spot or thermal drift degrades von Hamos resolution. Mitigation: ray-tracing gate on the apparent point/line focus, "
        "8 h stability acceptance test, water cooling and rigid source-sample mechanics.\n"
        "- A kW-class source exceeds the EUR 40k cap. Mitigation: request tube and generator prices separately, retain the "
        "100 kV/500 W integrated fallback, and design mechanics/cooling for a later kW upgrade."
    ),
}


copy2(SOURCE, OUTPUT)
doc = Document(OUTPUT)

to_defined_seen = 0
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text == "To be defined.":
        to_defined_seen += 1
        if to_defined_seen == 1:
            paragraph.text = REPLACEMENTS["To be defined."]
        elif to_defined_seen == 2:
            paragraph.text = (
                "The high-power source will be commissioned at the Mo, Rh, Ag and Cd K-edges, followed by at least one "
                "demonstration near 40 keV. XANES will be acquired scan-free over the HAPG bandwidth. Performance will be "
                "reported as acquisition time, SNR, energy resolution, photons per CZT channel, dead time and harmonic "
                "contamination. The Ag K-edge milestone is a quantitative spectrum in less than the 12 h benchmark; "
                "fluorescence measurements on a diluted or operando sample will demonstrate the advantage of CZT energy discrimination."
            )
        elif to_defined_seen == 3:
            paragraph.text = (
                "Results will be disseminated through open technical documentation of the von Hamos ray tracing, source "
                "acceptance protocol and CZT reconstruction, together with peer-reviewed publications and conference presentations."
            )
        continue

    for prefix, replacement in REPLACEMENTS.items():
        if prefix != "To be defined." and text.startswith(prefix):
            paragraph.text = replacement
            break

for paragraph in doc.paragraphs:
    if paragraph.text.strip().startswith(("WP1 -", "WP2 -", "WP3 -", "WP4 -", "WP5 -")):
        paragraph.paragraph_format.keep_with_next = True

doc.save(OUTPUT)
print(OUTPUT)
